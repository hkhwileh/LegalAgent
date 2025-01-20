try:
    from docx import Document
except ImportError:
    from docx.api import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import arabic_reshaper
from bidi.algorithm import get_display

class DocumentExporter:
    def __init__(self):
        # Register Arabic font for PDF
        try:
            pdfmetrics.registerFont(TTFont('Arabic', 'fonts/NotoNaskhArabic-Regular.ttf'))
        except:
            pass  # Fall back to default font if Arabic font is not available

    def export_to_pdf(self, content: dict) -> bytes:
        """Export the analysis results to PDF format."""
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        
        # Set font for Arabic text
        try:
            c.setFont('Arabic', 14)
        except:
            c.setFont('Helvetica', 14)

        y = 750  # Starting y position
        
        # Add title
        title = "تحليل المستند القانوني"
        title = get_display(arabic_reshaper.reshape(title))
        c.drawString(500, y, title)
        y -= 30

        # Add summary
        summary_title = "ملخص المستند"
        summary_title = get_display(arabic_reshaper.reshape(summary_title))
        c.drawString(500, y, summary_title)
        y -= 20
        
        summary_text = get_display(arabic_reshaper.reshape(content['summary']))
        # Wrap text to fit page width
        words = summary_text.split()
        line = ""
        for word in words:
            if c.stringWidth(line + word, 'Arabic', 12) < 500:
                line += word + " "
            else:
                c.drawString(50, y, line)
                y -= 15
                line = word + " "
        if line:
            c.drawString(50, y, line)
        y -= 30

        # Add legal analysis
        analysis_title = "تحليل المخالفات القانونية"
        analysis_title = get_display(arabic_reshaper.reshape(analysis_title))
        c.drawString(500, y, analysis_title)
        y -= 20
        
        analysis_text = get_display(arabic_reshaper.reshape(content['legal_analysis']))
        words = analysis_text.split()
        line = ""
        for word in words:
            if c.stringWidth(line + word, 'Arabic', 12) < 500:
                line += word + " "
            else:
                c.drawString(50, y, line)
                y -= 15
                line = word + " "
        if line:
            c.drawString(50, y, line)
        y -= 30

        # Add legislation mapping
        mapping_title = "الخريطة التشريعية"
        mapping_title = get_display(arabic_reshaper.reshape(mapping_title))
        c.drawString(500, y, mapping_title)
        y -= 20
        
        mapping_text = get_display(arabic_reshaper.reshape(content['legislation_mapping']))
        words = mapping_text.split()
        line = ""
        for word in words:
            if c.stringWidth(line + word, 'Arabic', 12) < 500:
                line += word + " "
            else:
                c.drawString(50, y, line)
                y -= 15
                line = word + " "
        if line:
            c.drawString(50, y, line)

        c.save()
        return buffer.getvalue()

    def export_to_word(self, content: dict) -> bytes:
        """Export the analysis results to Word format."""
        doc = Document()
        
        # Add title
        doc.add_heading("تحليل المستند القانوني", 0)
        
        # Add summary section
        doc.add_heading("ملخص المستند", level=1)
        doc.add_paragraph(content['summary'])
        
        # Add legal analysis section
        doc.add_heading("تحليل المخالفات القانونية", level=1)
        doc.add_paragraph(content['legal_analysis'])
        
        # Add legislation mapping section
        doc.add_heading("الخريطة التشريعية", level=1)
        doc.add_paragraph(content['legislation_mapping'])
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()